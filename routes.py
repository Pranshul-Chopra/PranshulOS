# ── routes.py ─────────────────────────────────────────────────────────────────
# Flask routes: Home, Dashboard, Docs.

import os
from flask import Blueprint, render_template, render_template_string, request, jsonify
from datetime import datetime as _dt, date as _date

import db
import insights as _insights

bp = Blueprint("main", __name__, template_folder="templates")

@bp.route("/")
def shell():
    return render_template("shell.html")



# ── Home ───────────────────────────────────────────────────────────────────────

_HOME_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <title>Home &mdash; PranshulOS</title>
  <link rel="stylesheet" href="/static/fonts/ibmflex.css"/>
  <style>
    *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
    :root {
      --bg: #111110; --surface: #1a1918; --surface2: #222120;
      --border: #2e2c2a; --border2: #3a3835;
      --text: #edeae4; --text2: #9b9690; --text3: #5c5955;
      --amber: #e8a84c; --amber-dim: #7a5820;
      --amber-faint: rgba(232,168,76,0.08); --amber-faint2: rgba(232,168,76,0.16);
      --green: #5aab7f;
      --radius: 10px;
      --mono: 'IBM Plex Mono', monospace; --sans: 'IBM Plex Sans', sans-serif;
    }
    #home-root {
      display: flex; flex-direction: column;
      width: 100%; height: 100%;
      background: var(--bg); color: var(--text);
      font-family: var(--sans); font-size: 15px; line-height: 1.6;
      -webkit-font-smoothing: antialiased;
    }
    .home-top {
      flex-shrink: 0; padding: 22px 28px 16px;
      border-bottom: 1px solid var(--border);
      display: flex; align-items: flex-start; justify-content: space-between; gap: 24px;
    }
    .greeting { font-size: 21px; font-weight: 400; letter-spacing: -0.01em; color: var(--text); }
    .greeting em { color: var(--amber); font-style: normal; }
    .insight-strip {
      display: flex; gap: 8px; flex-wrap: wrap; align-items: flex-start;
      padding-top: 4px;
    }
    .insight-pill {
      display: flex; align-items: center; gap: 6px;
      padding: 5px 11px; border-radius: 20px;
      border: 1px solid var(--border2); background: var(--surface);
      font-family: var(--mono); font-size: 11px; color: var(--text2);
      letter-spacing: 0.02em; line-height: 1.4; white-space: nowrap;
      animation: fadein 0.3s ease;
    }
    .insight-pill.streak { border-color: var(--amber-dim); color: var(--amber); background: var(--amber-faint); }
    .insight-pill.habit  { border-color: #3a6647; color: var(--green); background: rgba(90,171,127,0.07); }
    .insight-pill.late   { border-color: var(--border2); }
    .insight-pill.workload { border-color: var(--amber-dim); color: var(--text); background: var(--amber-faint); }
    .insight-pill .pill-icon { font-size: 13px; }
    @keyframes fadein { from { opacity: 0; transform: translateY(2px); } to { opacity: 1; } }
    .home-body {
      display: grid; grid-template-columns: 1fr 300px;
      flex: 1; min-height: 0; overflow: hidden;
    }
    .home-left {
      padding: 22px 24px 24px 28px;
      display: flex; flex-direction: column; gap: 22px;
      overflow-y: auto; border-right: 1px solid var(--border);
    }
    .home-left::-webkit-scrollbar { width: 3px; }
    .home-left::-webkit-scrollbar-track { background: transparent; }
    .home-left::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 2px; }
    .section-label {
      font-family: var(--mono); font-size: 11px; font-weight: 500;
      color: var(--text3); letter-spacing: 0.06em;
      text-transform: uppercase; margin-bottom: 12px;
    }
    .launch-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 8px; }
    .launch-btn {
      position: relative; padding: 16px 10px 14px; border-radius: var(--radius);
      background: var(--surface); border: 1px solid var(--border2);
      color: var(--text); font-family: var(--sans); text-align: center; cursor: pointer;
      transition: background 0.14s, border-color 0.14s, transform 0.1s;
      display: flex; flex-direction: column; align-items: center; gap: 8px;
    }
    .launch-btn:hover { background: var(--surface2); border-color: var(--amber-dim); transform: translateY(-1px); }
    .launch-btn .icon { font-size: 20px; line-height: 1; }
    .launch-btn .label { font-size: 12px; color: var(--text2); white-space: nowrap; overflow: hidden; text-overflow: ellipsis; max-width: 100%; }
    .launch-btn .del-x {
      position: absolute; top: 5px; right: 5px; width: 16px; height: 16px;
      border-radius: 50%; background: var(--surface2); border: 1px solid var(--border2);
      color: var(--text2); font-size: 10px; opacity: 0; transition: opacity 0.12s;
      display: flex; align-items: center; justify-content: center; z-index: 2; cursor: pointer;
    }
    .launch-btn:hover .del-x { opacity: 1; }
    .launch-btn .del-x:hover { background: #c0392b; border-color: #c0392b; color: #fff; }
    .add-btn {
      padding: 16px 10px 14px; border-radius: var(--radius);
      background: var(--surface); border: 1px dashed var(--border2);
      color: var(--text3); font-family: var(--sans); text-align: center; cursor: pointer;
      transition: background 0.14s, border-color 0.14s, transform 0.1s;
      display: flex; flex-direction: column; align-items: center; gap: 8px;
    }
    .add-btn:hover { background: var(--surface2); border-color: var(--amber-dim); transform: translateY(-1px); }
    .add-btn .add-bubble {
      width: 32px; height: 32px; border-radius: 50%; background: var(--amber);
      display: flex; align-items: center; justify-content: center; font-size: 18px; color: #111110;
    }
    .add-btn .label { font-size: 12px; color: var(--text3); }
    .cmd-bar {
      display: flex; gap: 8px; align-items: center;
      background: var(--surface); border: 1px solid var(--border2);
      border-radius: var(--radius); padding: 4px 4px 4px 14px; transition: border-color 0.15s;
    }
    .cmd-bar:focus-within { border-color: var(--amber-dim); }
    .cmd-input {
      flex: 1; background: transparent; border: none; outline: none;
      color: var(--text); font-family: var(--sans); font-size: 14px; padding: 7px 0;
    }
    .cmd-input::placeholder { color: var(--text3); }
    .cmd-go {
      padding: 7px 16px; border-radius: 7px; background: var(--amber); color: #111110;
      border: none; font-family: var(--sans); font-size: 13px; font-weight: 500; cursor: pointer;
    }
    .cmd-go:hover { opacity: 0.88; }
    .log {
      margin-top: 10px; padding: 11px 14px; border-radius: var(--radius);
      background: var(--surface); border: 1px solid var(--border);
      font-family: var(--mono); font-size: 13px; color: var(--text2); line-height: 1.8;
    }
    .log .entry.you { color: var(--text); }
    .log .entry.reply { color: var(--amber); }
    @keyframes fadein { from { opacity:0; transform: translateY(3px); } to { opacity:1; } }
    .log .entry { animation: fadein 0.18s ease; }
    .home-right {
      padding: 20px 16px; display: flex; flex-direction: column; gap: 12px; overflow-y: auto;
    }
    .home-right::-webkit-scrollbar { width: 3px; }
    .home-right::-webkit-scrollbar-track { background: transparent; }
    .home-right::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 2px; }
    .glance-card { background: var(--surface); border: 1px solid var(--border2); border-radius: 10px; overflow: hidden; display: flex; flex-direction: column; max-height: 260px; }
    .glance-header {
      display: flex; align-items: center; justify-content: space-between;
      padding: 11px 14px; border-bottom: 1px solid var(--border); flex-shrink: 0;
    }
    .glance-body-scroll { overflow-y: auto; flex: 1; min-height: 0; }
    .glance-body-scroll::-webkit-scrollbar { width: 3px; }
    .glance-body-scroll::-webkit-scrollbar-track { background: transparent; }
    .glance-body-scroll::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 2px; }
    .glance-title { font-family: var(--mono); font-size: 11px; letter-spacing: 0.07em; color: var(--amber); text-transform: uppercase; font-weight: 600; }
    .glance-count { font-family: var(--mono); font-size: 11px; color: var(--text2); }
    .glance-count .done { color: var(--green); }
    .glance-item { display: flex; align-items: center; gap: 10px; padding: 9px 14px; border-bottom: 1px solid var(--border); }
    .glance-item:last-child { border-bottom: none; }
    .g-dot { width: 6px; height: 6px; border-radius: 50%; flex-shrink: 0; }
    .g-dot.pending { background: transparent; border: 1.5px solid var(--text3); }
    .g-dot.done    { background: var(--green); border: 1.5px solid var(--green); }
    .g-text { flex: 1; font-size: 13px; color: var(--text); line-height: 1.5; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
    .g-time {
      font-family: var(--mono); font-size: 10px; color: var(--amber);
      background: var(--amber-faint); border: 1px solid var(--amber-dim);
      border-radius: 4px; padding: 1px 6px; letter-spacing: 0.03em;
      flex-shrink: 0; white-space: nowrap;
    }
    .glance-item.is-done .g-text { text-decoration: line-through; color: var(--text3); }
    .glance-subsection {
      padding: 6px 14px; font-family: var(--mono); font-size: 10px; letter-spacing: 0.07em;
      color: var(--text3); text-transform: uppercase;
      background: rgba(232,168,76,0.04); border-bottom: 1px solid var(--border); border-top: 1px solid var(--border);
    }
    .glance-empty { padding: 22px 14px; text-align: center; font-family: var(--mono); font-size: 12px; color: var(--text2); letter-spacing: 0.04em; line-height: 1.7; }
    .modal-overlay {
      position: fixed; inset: 0; background: rgba(0,0,0,0.65);
      display: flex; align-items: center; justify-content: center;
      z-index: 999; opacity: 0; pointer-events: none; transition: opacity 0.15s;
    }
    .modal-overlay.open { opacity: 1; pointer-events: all; }
    .modal { background: #1a1918; border: 1px solid var(--border2); border-radius: 14px; padding: 28px 28px 24px; width: 360px; max-width: calc(100vw - 48px); transform: translateY(8px); transition: transform 0.15s; }
    .modal-overlay.open .modal { transform: translateY(0); }
    .modal-title { font-size: 16px; font-weight: 500; color: var(--text); margin-bottom: 20px; }
    .field-label { font-family: var(--mono); font-size: 11px; color: var(--text3); letter-spacing: 0.05em; text-transform: uppercase; margin-bottom: 6px; }
    .field-group { margin-bottom: 14px; }
    .field-input { width: 100%; padding: 10px 13px; border-radius: 8px; background: var(--surface2); border: 1px solid var(--border2); color: var(--text); font-family: var(--sans); font-size: 14px; outline: none; transition: border-color 0.14s; }
    .field-input:focus { border-color: var(--amber-dim); }
    .field-input::placeholder { color: var(--text3); }
    .modal-actions { display: flex; gap: 8px; margin-top: 20px; }
    .modal-cancel { flex: 1; padding: 9px; border-radius: 8px; border: 1px solid var(--border2); background: transparent; color: var(--text2); font-family: var(--sans); font-size: 14px; cursor: pointer; }
    .modal-cancel:hover { background: var(--surface2); }
    .modal-save { flex: 2; padding: 9px; border-radius: 8px; border: none; background: var(--amber); color: #111110; font-family: var(--sans); font-size: 14px; font-weight: 500; cursor: pointer; }
    .modal-save:hover { opacity: 0.88; }
  </style>
</head>
<body>
<div id="home-root">

  <div class="home-top">
    <div class="greeting">{{ greeting }}</div>
    <div class="insight-strip" id="insight-strip">
      <span class="insight-pill" style="color:var(--text3)">Loading insights\u2026</span>
    </div>
  </div>

  <div class="home-body">

    <div class="home-left">
      <div>
        <div class="section-label">Quick launch</div>
        <div class="launch-grid" id="launch-grid">
          <button class="launch-btn" onclick="launch('youtube')"><span class="icon">\U0001f3ac</span><span class="label">YouTube</span></button>
          <button class="launch-btn" onclick="launch('spotify')"><span class="icon">\U0001f3b5</span><span class="label">Spotify</span></button>
          <button class="launch-btn" onclick="launch('whatsapp')"><span class="icon">\U0001f4ac</span><span class="label">WhatsApp</span></button>
          <button class="launch-btn" onclick="launch('discord')"><span class="icon">\U0001f3ae</span><span class="label">Discord</span></button>
          <button class="launch-btn" onclick="launch('github')"><span class="icon">\U0001f419</span><span class="label">GitHub</span></button>
          <button class="launch-btn" onclick="launch('linkedin')"><span class="icon">\U0001f517</span><span class="label">LinkedIn</span></button>
          <button class="launch-btn" onclick="launch('gmail')"><span class="icon">\u2709\ufe0f</span><span class="label">Gmail</span></button>
          <button class="launch-btn" onclick="launch('reddit')"><span class="icon">\U0001f47d</span><span class="label">Reddit</span></button>
          <button class="add-btn" id="add-btn" onclick="openModal()"><span class="add-bubble">+</span><span class="label">Add</span></button>
        </div>
      </div>
      <div>
        <div class="section-label">Or just tell me</div>
        <div class="cmd-bar">
          <input class="cmd-input" id="inp" placeholder='try "bored", "github", "spotify"\u2026' onkeydown="if(event.key==='Enter') go()"/>
          <button class="cmd-go" onclick="go()">Go</button>
        </div>
        <div class="log" id="log"><span style="color:var(--text3)">\u2192 what do you need?</span></div>
      </div>
    </div>

    <div class="home-right">
      <div class="glance-card" id="weekly-summary-card" style="max-height:none">
        <div class="glance-header">
          <span class="glance-title">\U0001f4c8 This Week</span>
          <span class="glance-count" id="weekly-trend-badge"></span>
        </div>
        <div id="weekly-summary-body" style="padding:12px 14px; font-family:var(--mono); font-size:12px; color:var(--text2); line-height:1.9;">
          <span style="color:var(--text3)">Loading\u2026</span>
        </div>
      </div>
      <div class="glance-card">
        <div class="glance-header">
          <span class="glance-title">\U0001f501 Routine</span>
          <span class="glance-count" id="routine-count"></span>
        </div>
        <div class="glance-body-scroll" id="routine-glance-body"></div>
      </div>
      <div class="glance-card">
        <div class="glance-header">
          <span class="glance-title">\U0001f4cb Tasks</span>
          <span class="glance-count" id="tasks-count"></span>
        </div>
        <div class="glance-body-scroll" id="tasks-glance-body"></div>
      </div>
    </div>

  </div>
</div>

<div class="modal-overlay" id="modal-overlay" onclick="overlayClick(event)">
  <div class="modal">
    <div class="modal-title">Add launcher</div>
    <div class="field-group"><div class="field-label">Name</div><input class="field-input" id="m-name" placeholder="e.g. Notion" maxlength="20"/></div>
    <div class="field-group"><div class="field-label">Icon</div><input class="field-input" id="m-icon" placeholder="\U0001f680" maxlength="4"/></div>
    <div class="field-group"><div class="field-label">URL</div><input class="field-input" id="m-url" placeholder="https://example.com"/></div>
    <div class="modal-actions">
      <button class="modal-cancel" onclick="closeModal()">Cancel</button>
      <button class="modal-save" onclick="saveCustom()">Add launcher</button>
    </div>
  </div>
</div>

<script>
function todayStr() {
  const d = new Date();
  return `${d.getFullYear()}-${String(d.getMonth()+1).padStart(2,'0')}-${String(d.getDate()).padStart(2,'0')}`;
}
async function launch(app) {
  try { await fetch('/launch/' + app); addLog('\u2192 opened ' + app, 'reply'); }
  catch(e) { addLog('\u2192 error', 'reply'); }
}
async function loadCustom() {
  try { renderCustom(await fetch('/api/launchers').then(r => r.json())); }
  catch(e) { console.error('loadCustom', e); }
}
function renderCustom(list) {
  const grid = document.getElementById('launch-grid');
  const add  = document.getElementById('add-btn');
  grid.querySelectorAll('.custom-btn').forEach(e => e.remove());
  list.forEach(item => {
    const btn = document.createElement('button');
    btn.className = 'launch-btn custom-btn';
    btn.innerHTML = `<span class="icon">${item.icon}</span><span class="label">${item.name}</span><span class="del-x">\xd7</span>`;
    btn.querySelector('.del-x').addEventListener('click', async e => {
      e.stopPropagation();
      await fetch('/api/launchers/' + item.id, {method:'DELETE'});
      addLog('\u2192 removed ' + item.name, 'reply');
      loadCustom();
    });
    btn.addEventListener('click', async e => {
      if (e.target.classList.contains('del-x')) return;
      try { await fetch('/launch/custom/' + item.id); addLog('\u2192 opened ' + item.name, 'reply'); }
      catch(e) { addLog('\u2192 error', 'reply'); }
    });
    grid.insertBefore(btn, add);
  });
}
async function loadRoutineGlance() {
  const body = document.getElementById('routine-glance-body');
  const cnt  = document.getElementById('routine-count');
  try {
    const today = todayStr();
    const [baseItems, weeklyItems] = await Promise.all([
      fetch('/api/routine?date=' + today).then(r => r.json()),
      fetch('/api/weekly-routine?date=' + today).then(r => r.json()),
    ]);
    // Tag source so checks go to the right endpoint
    baseItems.forEach(x => x._src = 'base');
    weeklyItems.forEach(x => x._src = 'weekly');
    const items = [...baseItems, ...weeklyItems];
    // sort timed first, chronological
    items.sort((a,b)=>{
      const at=a.time_start||null, bt=b.time_start||null;
      if(at&&bt) return at<bt?-1:at>bt?1:0;
      if(at) return -1; if(bt) return 1; return 0;
    });

    if (!items.length) {
      body.innerHTML = '<div class="glance-empty">NO ROUTINE ITEMS YET</div>';
      cnt.innerHTML = '';
      return;
    }
    const done = items.filter(x => x.checked).length;
    cnt.innerHTML = `<span class="done">${done}</span> / ${items.length}`;
    body.innerHTML = '';
    const fmtT = hhmm => {
      if (!hhmm) return '';
      const [h,m] = hhmm.split(':').map(Number);
      const ap = h >= 12 ? 'PM' : 'AM';
      return `${h%12||12}:${String(m).padStart(2,'0')} ${ap}`;
    };
    items.forEach(item => {
      const d = document.createElement('div');
      d.className = 'glance-item' + (item.checked ? ' is-done' : '');
      let timeHtml = '';
      if (item.time_start) {
        const end = item.time_end ? ' – ' + fmtT(item.time_end) : '';
        timeHtml = `<span class="g-time">${fmtT(item.time_start)}${end}</span>`;
      }
      d.innerHTML = `<span class="g-dot ${item.checked ? 'done' : 'pending'}"></span>${timeHtml}<span class="g-text">${item.text}</span>`;
      body.appendChild(d);
    });
  } catch(e) { body.innerHTML = '<div class="glance-empty">UNAVAILABLE</div>'; }
}
async function loadTasksGlance() {
  const body = document.getElementById('tasks-glance-body');
  const cnt  = document.getElementById('tasks-count');
  try {
    const data    = await fetch('/api/dashboard/state').then(r=>r.json());
    const today   = todayStr();
    const tasks   = data.tasks || [];
    const pending = tasks.filter(t=>!t.done && t.date===today);
    const done    = tasks.filter(t=> t.done && t.done_date===today);
    if (!pending.length && !done.length) { body.innerHTML='<div class="glance-empty">ALL CLEAR \u2014 NOTHING DUE TODAY</div>'; cnt.innerHTML=''; return; }
    cnt.innerHTML = `<span class="done">${done.length}</span> / ${pending.length+done.length}`;
    body.innerHTML = '';
    // sort timed tasks first, then untimed
    pending.sort((a,b)=>{
      const at=a.time_start||null, bt=b.time_start||null;
      if(at&&bt) return at<bt?-1:at>bt?1:0;
      if(at) return -1; if(bt) return 1; return 0;
    });
    pending.forEach(t => {
      const d = document.createElement('div');
      d.className='glance-item';
      let timeLabel='';
      if(t.time_start){
        const fmtT=hhmm=>{const[h,m]=hhmm.split(':').map(Number);const ap=h>=12?'PM':'AM';return`${h%12||12}:${String(m).padStart(2,'0')} ${ap}`;};
        timeLabel=`<span class="g-time">${fmtT(t.time_start)}${t.time_end?' – '+fmtT(t.time_end):''}</span>`;
      }
      d.innerHTML=`<span class="g-dot pending"></span>${timeLabel}<span class="g-text">${t.text}</span>`;
      body.appendChild(d);
    });
    if (done.length) {
      body.innerHTML += '<div class="glance-subsection">Done</div>';
      done.forEach(t => {
        const d = document.createElement('div');
        d.className='glance-item is-done';
        d.innerHTML=`<span class="g-dot done"></span><span class="g-text">${t.text}</span>`;
        body.appendChild(d);
      });
    }
  } catch(e) { body.innerHTML='<div class="glance-empty">UNAVAILABLE</div>'; }
}
function addLog(msg, cls) {
  const log = document.getElementById('log'); if(!log) return;
  const e = document.createElement('div');
  e.className='entry '+cls; e.textContent=msg;
  log.innerHTML=''; log.appendChild(e);
}
async function go() {
  const inp = document.getElementById('inp');
  const q = inp.value.trim(); if(!q) return;
  addLog('\u2192 '+q,'you'); inp.value='';
  try {
    const d = await fetch('/api/trigger',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({text:q})}).then(r=>r.json());
    addLog('\u2192 '+(d.reply||d.result||'ok'),'reply');
  } catch(e) { addLog('\u2192 error','reply'); }
}
function openModal()    { document.getElementById('modal-overlay').classList.add('open'); }
function closeModal()   { document.getElementById('modal-overlay').classList.remove('open'); }
function overlayClick(e){ if(e.target===document.getElementById('modal-overlay')) closeModal(); }
async function saveCustom() {
  const name = document.getElementById('m-name').value.trim();
  const icon = document.getElementById('m-icon').value.trim() || '\U0001f517';
  const url  = document.getElementById('m-url').value.trim();
  if (!name || !url) return;
  try {
    await fetch('/api/launchers', {method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({ name, icon, kind: 'url', target: url })});
    closeModal();
    ['m-name','m-icon','m-url'].forEach(id => document.getElementById(id).value = '');
    addLog('\u2192 added ' + name, 'reply');
    loadCustom();
  } catch(e) { addLog('\u2192 error saving launcher', 'reply'); }
}
// ── Insights ─────────────────────────────────────────────────────────────────

const PILL_CONFIG = {
  yesterday: { icon: '\u2713', cls: '' },
  streak:    { icon: '\U0001f525', cls: 'streak' },
  habit:     { icon: '\U0001f33f', cls: 'habit' },
  late:      { icon: '\U0001f319', cls: 'late' },
  workload:  { icon: '\u26a0\ufe0f', cls: 'workload' },
  trend_up:  { icon: '\u2197\ufe0f', cls: 'habit' },
  trend_down:{ icon: '\u2198\ufe0f', cls: 'late' },
};

async function loadInsights() {
  try {
    const data = await fetch('/api/insights').then(r => r.json());
    renderInsightStrip(data.greeting);
    renderWeeklySummary(data.weekly);
  } catch(e) {
    console.error('loadInsights', e);
    const strip = document.getElementById('insight-strip');
    if (strip) strip.innerHTML = '';
  }
}

function renderInsightStrip(greeting) {
  const strip = document.getElementById('insight-strip');
  if (!strip) return;
  const items = greeting && greeting.insights ? greeting.insights : [];
  if (!items.length) { strip.innerHTML = ''; return; }
  strip.innerHTML = items.map(ins => {
    const cfg = PILL_CONFIG[ins.type] || { icon: '\u2022', cls: '' };
    return `<span class="insight-pill ${cfg.cls}">
      <span class="pill-icon">${cfg.icon}</span>${ins.label}
    </span>`;
  }).join('');
}

function barBlocks(pct, len = 10) {
  const filled = Math.round((pct / 100) * len);
  return '\u2588'.repeat(filled) + '\u2591'.repeat(len - filled);
}

function renderWeeklySummary(weekly) {
  const body   = document.getElementById('weekly-summary-body');
  const badge  = document.getElementById('weekly-trend-badge');
  if (!body || !weekly || !weekly.days) return;

  const avg  = weekly.weekly_avg || 0;
  const best = weekly.best_day;
  const worst = weekly.worst_day;
  const trend = weekly.trend;

  // Trend badge
  if (badge && trend !== null && trend !== undefined) {
    const sign  = trend >= 0 ? '+' : '';
    const color = trend >= 0 ? 'var(--green)' : 'var(--text2)';
    badge.innerHTML = `<span style="color:${color};font-family:var(--mono);font-size:11px;">${sign}${Math.round(trend)}% vs last week</span>`;
  }

  // Bar chart
  let bars = weekly.days.map(d => {
    const bar = barBlocks(d.pct);
    return `<div style="display:flex;align-items:center;gap:8px;">
      <span style="width:26px;color:var(--text3)">${d.day}</span>
      <span style="color:var(--amber)">${bar}</span>
      <span style="color:var(--text3);font-size:10px">${d.pct}%</span>
    </div>`;
  }).join('');

  let meta = '';
  if (avg > 0) {
    meta += `<div style="margin-top:10px;padding-top:8px;border-top:1px solid var(--border);display:flex;gap:16px;font-size:11px;">`;
    meta += `<span style="color:var(--text2)">Avg <span style="color:var(--text)">${avg}%</span></span>`;
    if (best) meta += `<span style="color:var(--text2)">Best <span style="color:var(--green)">${best.day} (${best.pct}%)</span></span>`;
    if (worst && worst.day !== (best && best.day)) meta += `<span style="color:var(--text2)">Low <span style="color:var(--text)">${worst.day} (${worst.pct}%)</span></span>`;
    meta += `</div>`;
  }

  // Workload note
  let note = '';
  const tw = weekly.this_week_tasks, lw = weekly.last_week_tasks;
  if (tw && lw && lw > 0) {
    const delta = Math.round(((tw - lw) / lw) * 100);
    if (Math.abs(delta) >= 10) {
      const dir = delta > 0 ? '\u2191' : '\u2193';
      note = `<div style="margin-top:8px;font-size:11px;color:var(--text2)">${tw} tasks this week ${dir}${Math.abs(delta)}% vs last</div>`;
    }
  }

  body.innerHTML = bars + meta + note;
}

window.__pageInit = function() {
  const root = document.getElementById('home-root');
  if (root && root.parentElement) {
    root.parentElement.style.cssText = 'display:flex;flex-direction:column;flex:1;min-height:0;height:100%;';
  }
  loadCustom();
  loadRoutineGlance();
  loadTasksGlance();
  loadInsights();
};
window.launch=launch; window.go=go;
window.openModal=openModal; window.closeModal=closeModal;
window.overlayClick=overlayClick; window.saveCustom=saveCustom;
</script>
</body>
</html>"""


@bp.route("/home")
def home():
    h = _dt.now().hour
    if h < 12:   greeting = "Good morning \u2600\ufe0f"
    elif h < 17: greeting = "Good afternoon \U0001f324"
    else:        greeting = "Good evening \U0001f319"
    return render_template_string(_HOME_HTML, greeting=greeting)


# ── Dashboard ──────────────────────────────────────────────────────────────────

@bp.route("/dashboard")
def dashboard():
    return render_template("dashboard.html")


@bp.route("/api/dashboard/state", methods=["GET"])
def dashboard_state():
    today = _date.today().isoformat()
    db.rollover_tasks(today)
    return jsonify({"tasks": db.get_tasks(), "goals": db.get_goals()})


@bp.route("/api/dashboard/tasks", methods=["POST"])
def create_task():
    data       = request.get_json(silent=True) or {}
    text       = (data.get("text") or "").strip()
    date       = (data.get("date") or _date.today().isoformat()).strip()
    time_start = (data.get("time_start") or "").strip() or None
    time_end   = (data.get("time_end")   or "").strip() or None
    if not text:
        return jsonify({"error": "text required"}), 400
    return jsonify(db.add_task(text, date, time_start, time_end)), 201


@bp.route("/api/dashboard/tasks/<int:task_id>", methods=["PATCH"])
def patch_task(task_id):
    data = request.get_json(silent=True) or {}
    if "time_start" in data:
        data["time_start"] = (data.get("time_start") or "").strip() or None
    if "time_end" in data:
        data["time_end"] = (data.get("time_end") or "").strip() or None
    result = db.update_task(task_id, **data)
    if not result:
        return jsonify({"error": "not found"}), 404
    return jsonify(result)


@bp.route("/api/dashboard/tasks/<int:task_id>", methods=["DELETE"])
def remove_task(task_id):
    if db.delete_task(task_id):
        return jsonify({"ok": True})
    return jsonify({"error": "not found"}), 404


@bp.route("/api/dashboard/goals", methods=["POST"])
def create_goal():
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "text required"}), 400
    return jsonify(db.add_goal(text)), 201


@bp.route("/api/dashboard/goals/<int:goal_id>", methods=["DELETE"])
def remove_goal(goal_id):
    if db.delete_goal(goal_id):
        return jsonify({"ok": True})
    return jsonify({"error": "not found"}), 404


# ── Docs ───────────────────────────────────────────────────────────────────────

@bp.route("/docs")
def docs_page():
    return render_template("docs.html")


@bp.route("/api/docs", methods=["GET"])
def api_get_docs():
    return jsonify(db.get_all_docs())


@bp.route("/api/docs", methods=["POST"])
def api_create_doc():
    data  = request.get_json(silent=True) or {}
    title = data.get("title", "Untitled").strip() or "Untitled"
    return jsonify(db.create_doc(title)), 201


@bp.route("/api/docs/<int:doc_id>", methods=["GET"])
def api_get_doc(doc_id):
    doc = db.get_doc(doc_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    return jsonify(doc)


@bp.route("/api/docs/<int:doc_id>", methods=["PATCH"])
def api_update_doc(doc_id):
    data    = request.get_json(silent=True) or {}
    title   = data.get("title")
    content = data.get("content")
    pinned  = None
    if "pinned" in data:
        pinned = bool(data.get("pinned"))
    doc = db.update_doc(doc_id, title=title, content=content, pinned=pinned)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    return jsonify(doc)


@bp.route("/api/docs/<int:doc_id>", methods=["DELETE"])
def api_delete_doc(doc_id):
    if db.delete_doc(doc_id):
        return jsonify({"ok": True})
    return jsonify({"error": "Not found"}), 404


@bp.route("/api/docs/<int:doc_id>/export/<fmt>", methods=["GET"])
def api_export_doc(doc_id, fmt):
    """Export a document as .docx or .txt.
    GET /api/docs/<id>/export/docx  or  /api/docs/<id>/export/txt
    Returns the file directly so Electron / the browser triggers a native save dialog.
    """
    from flask import send_file
    import io

    fmt = (fmt or "").strip().lower()

    if fmt not in ("docx", "txt"):
        return jsonify({"error": f"Unsupported format: {fmt!r}. Use 'docx' or 'txt'."}), 400

    doc = db.get_doc(doc_id)
    if not doc:
        return jsonify({"error": "Not found"}), 404

    title   = doc.get("title") or "Untitled"
    content = doc.get("content") or ""

    if fmt == "txt":
        buf = io.BytesIO(content.encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            mimetype="text/plain; charset=utf-8",
            as_attachment=True,
            download_name=f"{title}.txt",
        )

    # --- .docx ---
    from docx import Document as _DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = _DocxDocument()

    # ── Minimal styling so the output looks intentional ──────────────────────
    style = document.styles["Normal"]
    font  = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title heading
    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Body — split on newlines; blank lines become paragraph breaks
    for line in content.split("\n"):
        para = document.add_paragraph(line)
        para.paragraph_format.space_after = Pt(0)

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{title}.docx",
    )


# ── Tickets ───────────────────────────────────────────────────────────────────

@bp.route("/tickets")
def tickets_page():
    return render_template("tickets.html")


@bp.route("/api/tickets", methods=["GET"])
def api_get_tickets():
    return jsonify(db.get_all_tickets())


@bp.route("/api/tickets", methods=["POST"])
def api_create_ticket():
    data = request.get_json(silent=True) or {}
    subject = (data.get("subject") or "").strip()
    if not subject:
        return jsonify({"error": "subject required"}), 400
    description = (data.get("description") or "").strip()
    priority = (data.get("priority") or "p2").strip().lower()
    due_at = (data.get("due_at") or "").strip() or None
    time_start = (data.get("time_start") or "").strip() or None
    time_end   = (data.get("time_end")   or "").strip() or None
    return jsonify(db.create_ticket(subject, description, priority, due_at, time_start, time_end)), 201


@bp.route("/api/tickets/<int:ticket_id>", methods=["PATCH"])
def api_update_ticket(ticket_id):
    data = request.get_json(silent=True) or {}
    result = db.update_ticket(ticket_id, **data)
    if not result:
        return jsonify({"error": "Not found"}), 404
    return jsonify(result)


@bp.route("/api/tickets/<int:ticket_id>", methods=["DELETE"])
def api_delete_ticket(ticket_id):
    if db.delete_ticket(ticket_id):
        return jsonify({"ok": True})
    return jsonify({"error": "Not found"}), 404


# ── Launch ─────────────────────────────────────────────────────────────────────
# Replaces the pywebview JS API (window.pywebview.api.open_X).
# HTML calls fetch('/launch/youtube') etc. Flask opens the URL in the
# system default browser via webbrowser.open().

import subprocess as _sp

# Suppress the console window flash on Windows for every subprocess call.
_CREATE_NO_WINDOW = 0x08000000 if os.name == "nt" else 0

def _open_url(url: str) -> None:
    """Open a URL in the user's existing default browser (new tab, not new process).
    Uses Windows 'start' shell command so the browser reuses its running instance
    instead of spawning a fresh one — avoids RAM spikes from webbrowser.open().
    CREATE_NO_WINDOW suppresses the terminal flash that previously appeared on
    every shortcut click."""
    try:
        _sp.Popen(
            ["cmd", "/c", "start", "", url],
            stdout=_sp.DEVNULL,
            stderr=_sp.DEVNULL,
            creationflags=_CREATE_NO_WINDOW,
        )
    except Exception as e:
        print(f"[launch] failed to open {url}: {e}")

_LAUNCH_URLS = {
    "youtube":       "https://youtube.com",
    "spotify":       "https://open.spotify.com",
    "whatsapp":      "https://web.whatsapp.com",
    "discord":       "https://discord.com/app",
    "github":        "https://github.com",
    "linkedin":      "https://www.linkedin.com",
    "gmail":         "https://mail.google.com",
    "reddit":        "https://www.reddit.com",
    "instagram":     "https://www.instagram.com",
    "twitch":        "https://twitch.tv",
    "roblox":        "https://www.roblox.com",
    "steam":         "https://store.steampowered.com",
    "warframe":      "https://warframe.market",
    "chatgpt":       "https://chatgpt.com",
    "drive":         "https://drive.google.com",
}

_TRIGGER_MAP = {
    "drive":         "drive",
    "reddit":        "reddit",
    "chatgpt":       "chatgpt",
    "open chatgpt":  "chatgpt",
    "youtube":       "youtube",
    "bored":         "youtube",
    "not feeling":   "youtube",
    "too lazy":      "youtube",
    "chill":         "youtube",
    "linkedin":      "linkedin",
    "github":        "github",
    "git":           "github",
    "whatsapp":      "whatsapp",
    "check messages":"whatsapp",
    "discord":       "discord",
    "twitch":        "twitch",
    "streaming":     "twitch",
    "roblox":        "roblox",
    "steam":         "steam",
    "play games":    "steam",
    "warframe":      "warframe",
    "spotify":       "spotify",
    "music":         "spotify",
    "instagram":     "instagram",
    "gmail":         "gmail",
    "open mail":     "gmail",
}


@bp.route("/launch/<app_name>")
def launch_app(app_name):
    url = _LAUNCH_URLS.get(app_name)
    if url:
        _open_url(url)
        return jsonify({"status": "ok", "opened": url})
    return jsonify({"status": "error", "message": "unknown app"}), 404


@bp.route("/api/trigger", methods=["POST"])
def trigger():
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").lower().strip()
    results, fired = [], set()
    for keyword, app_name in _TRIGGER_MAP.items():
        if keyword in text and app_name not in fired:
            url = _LAUNCH_URLS.get(app_name)
            if url:
                _open_url(url)
                results.append(f"Opened {app_name}")
                fired.add(app_name)
    return jsonify({
        "result": " · ".join(results) if results else None
    })


@bp.route("/api/launchers", methods=["GET"])
def get_launchers():
    return jsonify(db.get_launchers())


@bp.route("/api/launchers", methods=["POST"])
def create_launcher():
    data   = request.get_json(silent=True) or {}
    name   = (data.get("name") or "").strip()
    kind   = (data.get("kind") or "").strip()
    target = (data.get("target") or "").strip()
    icon   = (data.get("icon") or "🚀").strip()
    if not name or kind not in ("url", "path") or not target:
        return jsonify({"error": "name, kind (url|path), and target are required"}), 400
    return jsonify(db.add_launcher(name, kind, target, icon)), 201


@bp.route("/api/launchers/<int:launcher_id>", methods=["DELETE"])
def remove_launcher(launcher_id):
    if db.delete_launcher(launcher_id):
        return jsonify({"ok": True})
    return jsonify({"error": "not found"}), 404


@bp.route("/launch/custom/<int:launcher_id>")
def launch_custom(launcher_id):
    launchers = db.get_launchers()
    launcher  = next((l for l in launchers if l["id"] == launcher_id), None)
    if not launcher:
        return jsonify({"error": "not found"}), 404
    if launcher["kind"] == "url":
        _open_url(launcher["target"])
    else:
        try:
            _sp.Popen(
                ["cmd", "/c", "start", "", launcher["target"]],
                stdout=_sp.DEVNULL,
                stderr=_sp.DEVNULL,
                creationflags=_CREATE_NO_WINDOW,
            )
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    return jsonify({"status": "ok", "opened": launcher["target"]})


@bp.route("/api/ping")
def ping():
    """Electron polls this to know Flask is ready before opening the window."""
    return jsonify({"status": "ok"})


# ── Routine ────────────────────────────────────────────────────────────────────

@bp.route("/routine")
def routine_page():
    from flask import render_template_string
    return render_template_string(_ROUTINE_HTML)


@bp.route("/api/routine", methods=["GET"])
def api_get_routine():
    date = request.args.get("date") or _date.today().isoformat()
    return jsonify(db.get_routine_items(date))


@bp.route("/api/routine/items", methods=["POST"])
def api_add_routine_item():
    data       = request.get_json(silent=True) or {}
    text       = (data.get("text") or "").strip()
    time_start = (data.get("time_start") or "").strip() or None
    time_end   = (data.get("time_end")   or "").strip() or None
    if not text:
        return jsonify({"error": "text required"}), 400
    return jsonify(db.add_routine_item(text, time_start, time_end)), 201


@bp.route("/api/routine/items/<int:item_id>", methods=["PATCH"])
def api_patch_routine_item(item_id):
    data   = request.get_json(silent=True) or {}
    fields = {}
    if "text" in data:
        fields["text"] = (data.get("text") or "").strip()
    if "time_start" in data:
        fields["time_start"] = (data.get("time_start") or "").strip() or None
    if "time_end" in data:
        fields["time_end"] = (data.get("time_end") or "").strip() or None
    result = db.update_routine_item(item_id, **fields)
    if not result:
        return jsonify({"error": "not found"}), 404
    return jsonify(result)


@bp.route("/api/routine/items/<int:item_id>", methods=["DELETE"])
def api_delete_routine_item(item_id):
    if db.delete_routine_item(item_id):
        return jsonify({"ok": True})
    return jsonify({"error": "not found"}), 404


@bp.route("/api/routine/check", methods=["POST"])
def api_set_routine_check():
    data    = request.get_json(silent=True) or {}
    item_id = data.get("item_id")
    date    = (data.get("date") or _date.today().isoformat()).strip()
    checked = bool(data.get("checked", True))
    if not item_id:
        return jsonify({"error": "item_id required"}), 400
    db.set_routine_check(int(item_id), date, checked)
    return jsonify({"ok": True})


@bp.route("/api/routine/progress", methods=["GET"])
def api_routine_progress():
    date = request.args.get("date") or _date.today().isoformat()
    return jsonify(db.get_routine_progress(date))


@bp.route("/api/insights", methods=["GET"])
def api_insights():
    """Deterministic productivity insights — no AI required."""
    try:
        return jsonify(_insights.get_all_insights())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Weekly Routine API ─────────────────────────────────────────────────────────

@bp.route("/weekly-routine")
def weekly_routine_page():
    from flask import render_template_string
    return render_template_string(_WEEKLY_ROUTINE_HTML)


@bp.route("/api/weekly-routine", methods=["GET"])
def api_get_weekly_routine():
    """?date=YYYY-MM-DD  → items for that date's weekday (used by home glance)."""
    date = request.args.get("date") or _date.today().isoformat()
    return jsonify(db.get_weekly_items_for_date(date))


@bp.route("/api/weekly-routine/all", methods=["GET"])
def api_get_all_weekly():
    """All weekly items across all days (used by planner page)."""
    return jsonify(db.get_all_weekly_items())


@bp.route("/api/weekly-routine/items", methods=["POST"])
def api_add_weekly_item():
    data       = request.get_json(silent=True) or {}
    text       = (data.get("text") or "").strip()
    weekday    = data.get("weekday")
    time_start = (data.get("time_start") or "").strip() or None
    time_end   = (data.get("time_end")   or "").strip() or None
    if not text:
        return jsonify({"error": "text required"}), 400
    if weekday is None or int(weekday) not in range(7):
        return jsonify({"error": "weekday 0-6 required"}), 400
    return jsonify(db.add_weekly_item(text, int(weekday), time_start, time_end)), 201


@bp.route("/api/weekly-routine/items/<int:item_id>", methods=["PATCH"])
def api_patch_weekly_item(item_id):
    data   = request.get_json(silent=True) or {}
    fields = {}
    if "text" in data:
        fields["text"] = (data.get("text") or "").strip()
    if "time_start" in data:
        fields["time_start"] = (data.get("time_start") or "").strip() or None
    if "time_end" in data:
        fields["time_end"] = (data.get("time_end") or "").strip() or None
    result = db.update_weekly_item(item_id, **fields)
    if not result:
        return jsonify({"error": "not found"}), 404
    return jsonify(result)


@bp.route("/api/weekly-routine/items/<int:item_id>", methods=["DELETE"])
def api_delete_weekly_item(item_id):
    if db.delete_weekly_item(item_id):
        return jsonify({"ok": True})
    return jsonify({"error": "not found"}), 404


@bp.route("/api/weekly-routine/check", methods=["POST"])
def api_set_weekly_check():
    data    = request.get_json(silent=True) or {}
    item_id = data.get("item_id")
    date    = (data.get("date") or _date.today().isoformat()).strip()
    checked = bool(data.get("checked", True))
    if not item_id:
        return jsonify({"error": "item_id required"}), 400
    db.set_weekly_check(int(item_id), date, checked)
    return jsonify({"ok": True})


# ── Routine page template ──────────────────────────────────────────────────────

_ROUTINE_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <title>Routine — PranshulOS</title>
  <link rel="stylesheet" href="/static/fonts/ibmflex.css"/>
  <style>
    *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
    :root {
      --bg: #111110; --surface: #1a1918; --surface2: #222120;
      --border: #2e2c2a; --border2: #3a3835;
      --text: #edeae4; --text2: #9b9690; --text3: #5c5955;
      --amber: #e8a84c; --amber-dim: #7a5820;
      --amber-faint: rgba(232,168,76,0.08); --amber-faint2: rgba(232,168,76,0.16);
      --green: #5aab7f; --green-faint: rgba(90,171,127,0.10);
      --mono: 'IBM Plex Mono', monospace; --sans: 'IBM Plex Sans', sans-serif;
      --r: 10px;
    }
    html, body { height: 100%; background: var(--bg); color: var(--text);
      font-family: var(--sans); -webkit-font-smoothing: antialiased; }

    /* ── two-column shell ── */
    .routine-shell {
      display: grid;
      grid-template-columns: 1fr 300px;
      height: 100%;
      overflow: hidden;
    }
    .routine-left {
      padding: 36px 28px 60px;
      overflow-y: auto;
      border-right: 1px solid var(--border);
    }
    .routine-left::-webkit-scrollbar { width: 3px; }
    .routine-left::-webkit-scrollbar-track { background: transparent; }
    .routine-left::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 2px; }
    .routine-right {
      padding: 22px 18px;
      overflow-y: auto;
      display: flex;
      flex-direction: column;
      gap: 14px;
    }
    .routine-right::-webkit-scrollbar { width: 3px; }
    .routine-right::-webkit-scrollbar-track { background: transparent; }
    .routine-right::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 2px; }

    /* keep old .page for max-width on left col */
    .page { max-width: 560px; }

    /* ── header ── */
    .page-header { margin-bottom: 30px; }
    .page-title {
      font-family: var(--mono); font-size: 11px; letter-spacing: 0.08em;
      color: var(--text3); text-transform: uppercase; margin-bottom: 6px;
    }
    .page-sub { font-size: 22px; font-weight: 400; color: var(--text); letter-spacing: -0.01em; }
    .page-date { font-family: var(--mono); font-size: 12px; color: var(--text2); margin-top: 4px; }

    /* ── progress bar ── */
    .progress-wrap { margin-bottom: 28px; }
    .progress-label {
      display: flex; justify-content: space-between;
      font-family: var(--mono); font-size: 11px; color: var(--text3);
      letter-spacing: 0.05em; margin-bottom: 7px;
    }
    .progress-label .done-label { color: var(--green); }
    .progress-bar { height: 3px; background: var(--border2); border-radius: 2px; overflow: hidden; }
    .progress-fill { height: 100%; background: var(--green); border-radius: 2px; transition: width 0.35s ease; }

    /* ── add row ── */
    .add-row { display: flex; gap: 8px; margin-bottom: 20px; }
    .add-input {
      flex: 1; background: var(--surface); border: 1px solid var(--border2);
      border-radius: 8px; padding: 10px 13px;
      font-family: var(--sans); font-size: 14px; color: var(--text);
      outline: none; transition: border-color 0.15s;
    }
    .add-input:focus { border-color: var(--amber-dim); }
    .add-input::placeholder { color: var(--text3); }
    .add-btn-primary {
      background: var(--amber-faint2); border: 1px solid var(--amber-dim);
      border-radius: 8px; padding: 10px 16px;
      font-family: var(--mono); font-size: 11px; color: var(--amber);
      cursor: pointer; letter-spacing: 0.05em; white-space: nowrap;
      transition: background 0.15s;
    }
    .add-btn-primary:hover { background: rgba(232,168,76,0.24); }

    /* time slot row */
    .time-slot-row {
      display: flex; gap: 6px; align-items: center; margin-top: 8px; margin-bottom: 4px;
    }
    .time-slot-label {
      font-family: var(--mono); font-size: 10px; color: var(--text3);
      letter-spacing: 0.05em; flex-shrink: 0;
    }
    .time-inp {
      background: var(--surface); border: 1px solid var(--border2);
      border-radius: 6px; padding: 5px 8px;
      font-family: var(--mono); font-size: 12px; color: var(--text);
      outline: none; transition: border-color 0.15s; width: 96px;
    }
    .time-inp:focus { border-color: var(--amber-dim); }
    .time-inp::-webkit-calendar-picker-indicator { filter: invert(0.5); cursor: pointer; }
    .time-sep { font-family: var(--mono); font-size: 11px; color: var(--text3); }
    .time-clear {
      background: none; border: none; color: var(--text3);
      font-size: 13px; cursor: pointer; padding: 2px 5px; border-radius: 4px;
    }
    .time-clear:hover { color: #c87070; }

    /* inline time edit */
    .ri-edit-row {
      display: none; gap: 6px; align-items: center; padding: 8px 14px 10px 43px;
      background: var(--surface2); border: 1px solid var(--border);
      border-top: none; border-radius: 0 0 8px 8px; margin-top: -6px; margin-bottom: 6px;
    }
    .ri-edit-row.open { display: flex; }

    /* ── section labels ── */
    .sec-label {
      font-family: var(--mono); font-size: 11px; letter-spacing: 0.07em;
      color: var(--text3); text-transform: uppercase; margin-bottom: 10px;
    }

    /* ── routine list ── */
    .routine-list { display: flex; flex-direction: column; gap: 6px; margin-bottom: 28px; }
    .routine-item {
      display: flex; align-items: center; gap: 11px;
      padding: 12px 14px;
      background: var(--surface); border: 1px solid var(--border);
      border-radius: 8px; transition: border-color 0.15s, background 0.15s;
    }
    .routine-item.checked { background: var(--green-faint); border-color: rgba(90,171,127,0.2); }
    .routine-item.checked .ri-text { color: var(--text3); text-decoration: line-through; }

    .ri-check {
      width: 18px; height: 18px; border-radius: 50%;
      border: 1.5px solid var(--border2); flex-shrink: 0; cursor: pointer;
      display: flex; align-items: center; justify-content: center; transition: all 0.15s;
    }
    .ri-check.checked { border-color: var(--green); background: var(--green-faint); }
    .ri-check.checked::after {
      content:''; width: 7px; height: 5px;
      border-left: 1.5px solid var(--green); border-bottom: 1.5px solid var(--green);
      transform: rotate(-45deg) translateY(-1px); display: block;
    }
    .ri-text { flex: 1; font-size: 14px; color: var(--text); line-height: 1.4; }
    .ri-time {
      font-family: var(--mono); font-size: 10px; color: var(--amber);
      background: var(--amber-faint); border: 1px solid var(--amber-dim);
      border-radius: 4px; padding: 2px 7px; letter-spacing: 0.03em;
      flex-shrink: 0; white-space: nowrap; cursor: pointer;
    }
    .ri-time:hover { background: var(--amber-faint2); }
    .ri-time-add {
      font-family: var(--mono); font-size: 10px; color: var(--text3);
      background: none; border: 1px dashed var(--border2); border-radius: 4px;
      padding: 2px 7px; cursor: pointer; flex-shrink: 0; opacity: 0;
      transition: opacity 0.1s;
    }
    .routine-item:hover .ri-time-add { opacity: 1; }
    .ri-time-add:hover { color: var(--amber); border-color: var(--amber-dim); }
    .ri-del {
      background: none; border: none; color: transparent; cursor: pointer;
      font-size: 16px; padding: 2px 4px; border-radius: 4px;
      transition: color 0.1s; flex-shrink: 0;
    }
    .routine-item:hover .ri-del { color: var(--text3); }
    .ri-del:hover { color: #c87070 !important; }

    /* ── empty state ── */
    .empty {
      padding: 28px 12px; text-align: center;
      font-family: var(--mono); font-size: 12px; color: var(--text3);
      letter-spacing: 0.05em; border: 1px dashed var(--border); border-radius: 8px;
    }

    /* ── all-done card ── */
    .all-done {
      display: none; padding: 16px 18px;
      background: var(--green-faint); border: 1px solid rgba(90,171,127,0.25);
      border-radius: 8px; font-family: var(--mono); font-size: 12px; color: var(--green);
      letter-spacing: 0.05em; text-align: center; margin-top: -10px; margin-bottom: 20px;
    }
    .all-done.visible { display: block; }

    /* ── template hint ── */
    .template-hint {
      font-family: var(--mono); font-size: 11px; color: var(--text3);
      letter-spacing: 0.04em; margin-top: 6px; line-height: 1.8;
    }

    /* ── routine health ── */
    .health-panel-title {
      font-family: var(--mono); font-size: 10px; letter-spacing: 0.1em;
      color: var(--text3); text-transform: uppercase; margin-bottom: 10px;
    }
    .health-grid { display: flex; flex-direction: column; gap: 8px; }
    .health-card {
      background: var(--surface); border: 1px solid var(--border);
      border-radius: 10px; padding: 14px 16px; display: flex; flex-direction: column; gap: 4px;
    }
    .health-card.Healthy   { border-color: rgba(90,171,127,0.3); background: var(--green-faint); }
    .health-card.Growing   { border-color: rgba(90,171,127,0.2); }
    .health-card.Declining { border-color: rgba(200,112,112,0.25); }
    .health-card.Recovering{ border-color: var(--border2); }
    .hc-name { font-size: 13px; color: var(--text); line-height: 1.3; font-weight: 400; }
    .hc-state {
      font-family: var(--mono); font-size: 10px; letter-spacing: 0.06em;
      text-transform: uppercase; font-weight: 500; margin-top: 2px;
      display: flex; align-items: center; gap: 5px;
    }
    .hc-state.Healthy   { color: var(--green); }
    .hc-state.Growing   { color: #7ecfa3; }
    .hc-state.Stable    { color: var(--text2); }
    .hc-state.Recovering{ color: var(--amber); }
    .hc-state.Declining { color: #c87070; }
    .hc-meta { font-family: var(--mono); font-size: 10px; color: var(--text3); margin-top: 6px; line-height: 1.6; }
    .health-empty {
      font-family: var(--mono); font-size: 11px; color: var(--text3);
      letter-spacing: 0.04em; padding: 20px 0; text-align: center;
    }
  </style>
</head>
<body>
<div class="routine-shell">

  <!-- ── Left: checklist + add form ── -->
  <div class="routine-left">
  <div class="page">
  <div class="page-header">
    <div class="page-title">Daily Routine</div>
    <div class="page-sub">Routine Checklist</div>
    <div class="page-date" id="r-date"></div>
  </div>

  <div class="progress-wrap" id="progress-wrap" style="display:none">
    <div class="progress-label">
      <span>DAILY PROGRESS</span>
      <span class="done-label" id="progress-text">0 / 0</span>
    </div>
    <div class="progress-bar"><div class="progress-fill" id="progress-fill" style="width:0%"></div></div>
  </div>

  <div class="all-done" id="all-done">✓ ALL DONE — GREAT WORK TODAY</div>

  <div class="sec-label" id="sec-checklist" style="display:none">TODAY'S CHECKLIST</div>
  <div class="routine-list" id="routine-list"></div>

  <div class="sec-label">ADD ROUTINE ITEM</div>
  <div class="add-row">
    <input class="add-input" id="r-inp" placeholder="e.g. Morning stretch, Read 20 pages…"
      onkeydown="if(event.key==='Enter') addItem()"/>
    <button class="add-btn-primary" onclick="addItem()">+ ADD</button>
  </div>
  <div class="time-slot-row">
    <span class="time-slot-label">TIME SLOT</span>
    <input class="time-inp" type="time" id="r-ts-start" title="Start time">
    <span class="time-sep">–</span>
    <input class="time-inp" type="time" id="r-ts-end" title="End time">
    <button class="time-clear" title="Clear time slot" onclick="clearRoutineTime()">×</button>
  </div>
  <div class="template-hint">
    Items added here become part of your daily template — they appear every day
    with fresh checkboxes. Checks reset at midnight; the template never does.
  </div>
  </div><!-- .page -->
  </div><!-- .routine-left -->

  <!-- ── Right: routine health ── -->
  <div class="routine-right">
    <div class="health-panel-title">ROUTINE HEALTH</div>
    <div class="health-grid" id="health-grid">
      <div class="health-empty">Loading\u2026</div>
    </div>
  </div>

</div><!-- .routine-shell -->

<script>
let _items = [];
let _today = '';
let _editingKey = null;

function itemKey(item) { return item._src + ':' + item.id; }
function findItem(id, src) { return _items.find(x => x.id === id && x._src === src); }

function todayStr() {
  const d = new Date();
  return `${d.getFullYear()}-${String(d.getMonth()+1).padStart(2,'0')}-${String(d.getDate()).padStart(2,'0')}`;
}

function fmtToday() {
  return new Date().toLocaleDateString('en-US',
    { weekday:'long', month:'long', day:'numeric' });
}

function fmtTime(hhmm) {
  if (!hhmm) return '';
  const [h, m] = hhmm.split(':').map(Number);
  const ap = h >= 12 ? 'PM' : 'AM';
  return `${h % 12 || 12}:${String(m).padStart(2,'0')} ${ap}`;
}

function sortByTime(arr) {
  return [...arr].sort((a, b) => {
    const at = a.time_start || null, bt = b.time_start || null;
    if (at && bt) return at < bt ? -1 : at > bt ? 1 : 0;
    if (at) return -1;
    if (bt) return 1;
    return 0;
  });
}

function readRoutineTime() {
  const s = (document.getElementById('r-ts-start').value || '').trim();
  const e = (document.getElementById('r-ts-end').value   || '').trim();
  return { time_start: s || null, time_end: e || null };
}

function clearRoutineTime() {
  document.getElementById('r-ts-start').value = '';
  document.getElementById('r-ts-end').value   = '';
}

async function api(method, path, body) {
  const opts = { method, headers:{'Content-Type':'application/json'} };
  if (body) opts.body = JSON.stringify(body);
  const r = await fetch(path, opts);
  return r.json();
}

async function load() {
  _today = todayStr();
  document.getElementById('r-date').textContent = fmtToday();
  try {
    const [base, weekly] = await Promise.all([
      api('GET', '/api/routine?date=' + _today),
      api('GET', '/api/weekly-routine?date=' + _today),
    ]);
    base.forEach(x => x._src = 'base');
    weekly.forEach(x => x._src = 'weekly');
    _items = [...base, ...weekly];
  } catch(e) { _items = []; }
  render();
}

async function addItem() {
  const inp = document.getElementById('r-inp');
  const text = inp.value.trim();
  if (!text) return;
  const slot = readRoutineTime();
  try {
    const item = await api('POST', '/api/routine/items', { text, ...slot });
    item.checked = 0;
    item._src = 'base';
    _items.push(item);
    inp.value = '';
    clearRoutineTime();
    render();
  } catch(e) { console.error('addItem', e); }
}

async function saveItemTime(id, src, startId, endId) {
  const s = (document.getElementById(startId).value || '').trim() || null;
  const e = (document.getElementById(endId).value   || '').trim() || null;
  const patchPath = src === 'weekly'
    ? `/api/weekly-routine/items/${id}`
    : `/api/routine/items/${id}`;
  try {
    const updated = await api('PATCH', patchPath, { time_start: s, time_end: e });
    const idx = _items.findIndex(x => x.id === id && x._src === src);
    if (idx !== -1) _items[idx] = { ..._items[idx], ...updated };
    _editingKey = null;
    render();
  } catch(e) { console.error('saveItemTime', e); }
}

function toggleEditTime(id, src) {
  const key = src + ':' + id;
  _editingKey = (_editingKey === key) ? null : key;
  render();
}

async function toggleItem(id, src) {
  const item = findItem(id, src);
  if (!item) return;
  const newChecked = !item.checked;
  item.checked = newChecked ? 1 : 0;
  render();
  const endpoint = src === 'weekly' ? '/api/weekly-routine/check' : '/api/routine/check';
  try {
    await api('POST', endpoint, { item_id: id, date: _today, checked: newChecked });
  } catch(e) {
    // revert on failure
    item.checked = newChecked ? 0 : 1;
    render();
  }
}

async function deleteItem(id, src) {
  if (src !== 'base') return;
  _items = _items.filter(x => !(x.id === id && x._src === src));
  render();
  try {
    await api('DELETE', `/api/routine/items/${id}`);
  } catch(e) { console.error('deleteItem', e); }
}

function render() {
  const list = document.getElementById('routine-list');
  const wrap = document.getElementById('progress-wrap');
  const fill = document.getElementById('progress-fill');
  const ptext = document.getElementById('progress-text');
  const secLabel = document.getElementById('sec-checklist');
  const allDone = document.getElementById('all-done');
  if (!list) return;

  list.innerHTML = '';

  const total = _items.length;
  const done  = _items.filter(x => x.checked).length;

  if (total === 0) {
    const em = document.createElement('div');
    em.className = 'empty';
    em.textContent = 'NO ITEMS YET — ADD YOUR FIRST ROUTINE ITEM BELOW';
    list.appendChild(em);
    wrap.style.display = 'none';
    secLabel.style.display = 'none';
    allDone.classList.remove('visible');
  } else {
    wrap.style.display = '';
    secLabel.style.display = '';
    const pct = total > 0 ? Math.round((done / total) * 100) : 0;
    fill.style.width = pct + '%';
    ptext.textContent = `${done} / ${total}`;
    allDone.classList.toggle('visible', done === total && total > 0);

    sortByTime(_items).forEach(item => {
      const div = document.createElement('div');
      div.className = 'routine-item' + (item.checked ? ' checked' : '');

      const chk = document.createElement('div');
      chk.className = 'ri-check' + (item.checked ? ' checked' : '');
      chk.onclick = () => toggleItem(item.id, item._src);
      div.appendChild(chk);

      if (item.time_start) {
        const timeBadge = document.createElement('span');
        timeBadge.className = 'ri-time';
        const end = item.time_end ? ` – ${fmtTime(item.time_end)}` : '';
        timeBadge.textContent = fmtTime(item.time_start) + end;
        timeBadge.title = 'Click to edit time slot';
        timeBadge.onclick = () => toggleEditTime(item.id, item._src);
        div.appendChild(timeBadge);
      } else {
        const addTime = document.createElement('button');
        addTime.className = 'ri-time-add';
        addTime.textContent = '+ time';
        addTime.onclick = () => toggleEditTime(item.id, item._src);
        div.appendChild(addTime);
      }

      const txt = document.createElement('div');
      txt.className = 'ri-text';
      txt.textContent = item.text;
      div.appendChild(txt);

      if (item._src === 'base') {
        const del = document.createElement('button');
        del.className = 'ri-del';
        del.textContent = '×';
        del.title = 'Remove from template';
        del.onclick = () => deleteItem(item.id, item._src);
        div.appendChild(del);
      }

      list.appendChild(div);

      const key = itemKey(item);
      if (_editingKey === key) {
        const editRow = document.createElement('div');
        editRow.className = 'ri-edit-row open';
        const sid = 'edit-start-' + item._src + '-' + item.id;
        const eid = 'edit-end-' + item._src + '-' + item.id;
        editRow.innerHTML = `
          <span class="time-slot-label">TIME</span>
          <input class="time-inp" type="time" id="${sid}" value="${item.time_start || ''}">
          <span class="time-sep">–</span>
          <input class="time-inp" type="time" id="${eid}" value="${item.time_end || ''}">
          <button class="add-btn-primary" style="padding:6px 12px" onclick="saveItemTime(${item.id},'${item._src}','${sid}','${eid}')">SAVE</button>
          <button class="time-clear" id="clear-${item._src}-${item.id}">CLEAR</button>`;
        list.appendChild(editRow);
        const patchPath = item._src === 'weekly'
          ? `/api/weekly-routine/items/${item.id}`
          : `/api/routine/items/${item.id}`;
        editRow.querySelector('#clear-' + item._src + '-' + item.id).onclick = async () => {
          await api('PATCH', patchPath, { time_start: null, time_end: null });
          const idx = _items.findIndex(x => x.id === item.id && x._src === item._src);
          if (idx !== -1) { _items[idx].time_start = null; _items[idx].time_end = null; }
          _editingKey = null;
          render();
        };
      }
    });
  }
}

// ── Routine Health ────────────────────────────────────────────────────────────

async function loadRoutineHealth() {
  const grid = document.getElementById('health-grid');
  if (!grid) return;
  try {
    const data = await fetch('/api/insights').then(r => r.json());
    const items = data.routine_health || [];
    if (!items.length) {
      grid.innerHTML = '<div class="health-empty">Add routine items to see health data.</div>';
      return;
    }
    const STATE_ICON = { Healthy: '\U0001f7e2', Growing: '\U0001f4c8', Stable: '\u2796', Recovering: '\U0001f504', Declining: '\U0001f534' };
    grid.innerHTML = items.map(h => `
      <div class="health-card ${h.state}">
        <div class="hc-name">${h.text}</div>
        <div class="hc-state ${h.state}">${(STATE_ICON[h.state] || '')} ${h.state}</div>
        <div class="hc-meta">${h.streak > 0 ? h.streak + ' day streak' : 'No current streak'}<br>${h.completion_pct}% (30d)</div>
      </div>`).join('');
  } catch(e) {
    console.error('loadRoutineHealth', e);
    if (grid) grid.innerHTML = '<div class="health-empty">Could not load health data.</div>';
  }
}

// ── Expose to window ──────────────────────────────────────────────────────────
window.addItem         = addItem;
window.clearRoutineTime = clearRoutineTime;
window.saveItemTime    = saveItemTime;
window.__pageInit = function() {
  const shell = document.querySelector('.routine-shell');
  if (shell && shell.parentElement) {
    shell.parentElement.style.cssText = 'display:flex;flex-direction:column;flex:1;min-height:0;height:100%;';
  }
  load();
  loadRoutineHealth();
};
</script>
</body>
</html>"""


# ── Weekly Routine page ────────────────────────────────────────────────────────

_WEEKLY_ROUTINE_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <title>Weekly Routine \u2014 PranshulOS</title>
  <link rel="stylesheet" href="/static/fonts/ibmflex.css"/>
  <style>
    *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
    :root {
      --bg: #111110; --surface: #1a1918; --surface2: #222120;
      --border: #2e2c2a; --border2: #3a3835;
      --text: #edeae4; --text2: #9b9690; --text3: #5c5955;
      --amber: #e8a84c; --amber-dim: #7a5820;
      --amber-faint: rgba(232,168,76,0.08); --amber-faint2: rgba(232,168,76,0.16);
      --green: #5aab7f; --green-faint: rgba(90,171,127,0.10);
      --mono: 'IBM Plex Mono', monospace; --sans: 'IBM Plex Sans', sans-serif;
    }
    html, body { height: 100%; background: var(--bg); color: var(--text);
      font-family: var(--sans); -webkit-font-smoothing: antialiased; }
    body::-webkit-scrollbar { width: 3px; }
    body::-webkit-scrollbar-track { background: transparent; }
    body::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 2px; }
    .page { max-width: 1180px; margin: 0 auto; padding: 36px 28px 60px; max-height: 100vh; overflow-y: auto; scrollbar-width: thin; scrollbar-color: var(--border2) transparent; }
    .page::-webkit-scrollbar { width: 3px; height: 3px; }
    .page::-webkit-scrollbar-track { background: transparent; }
    .page::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 999px; }

    .page-head {
      display: flex; justify-content: space-between; align-items: flex-end; gap: 16px;
      margin-bottom: 24px;
    }
    .page-title {
      font-family: var(--mono); font-size: 11px; letter-spacing: 0.08em;
      color: var(--text3); text-transform: uppercase; margin-bottom: 6px;
    }
    .page-sub { font-size: 22px; font-weight: 400; letter-spacing: -0.01em; }
    .page-pill {
      font-family: var(--mono); font-size: 10px; letter-spacing: 0.06em;
      color: var(--amber); background: var(--amber-faint); border: 1px solid var(--amber-dim);
      border-radius: 999px; padding: 6px 10px; white-space: nowrap;
    }

    /* add bar */
    .add-bar {
      display: flex; gap: 8px; margin-bottom: 24px; align-items: center;
      background: var(--surface); border: 1px solid var(--border2);
      border-radius: 10px; padding: 6px 8px 6px 12px; flex-wrap: nowrap;
      max-width: 920px;
    }
    .add-bar:focus-within { border-color: var(--amber-dim); }
    .add-inp {
      flex: 1; max-width: 520px; background: transparent; border: none; outline: none;
      color: var(--text); font-family: var(--sans); font-size: 13px; padding: 4px 0;
    }
    .add-inp::placeholder { color: var(--text3); }
    .day-tabs { display: flex; gap: 6px; flex-shrink: 0; margin-left: 6px; }
    .day-tab {
      padding: 4px 8px; border-radius: 6px; border: 1px solid var(--border2);
      background: var(--bg); color: var(--text2);
      font-family: var(--mono); font-size: 10px; cursor: pointer;
      transition: all 0.12s; letter-spacing: 0.03em;
    }
    .day-tab.sel { background: var(--amber-faint2); border-color: var(--amber-dim); color: var(--amber); }
    .add-go {
      padding: 6px 10px; border-radius: 7px;
      background: var(--amber-faint2); border: 1px solid var(--amber-dim);
      color: var(--amber); font-family: var(--mono); font-size: 11px;
      cursor: pointer; letter-spacing: 0.04em; white-space: nowrap;
      transition: background 0.13s;
    }
    .add-go:hover { background: rgba(232,168,76,0.26); }

    /* 7-day grid */
    .planner-scroll {
      overflow-x: auto;
      overflow-y: hidden;
      padding-bottom: 8px;
      margin-bottom: 16px;
    }
    .planner-scroll::-webkit-scrollbar { height: 4px; width: 4px; }
    .planner-scroll::-webkit-scrollbar-track { background: transparent; }
    .planner-scroll::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 999px; }
    .week-grid { display: grid; grid-template-columns: repeat(7, minmax(140px, 1fr)); gap: 10px; min-width: 980px; }

    .day-col {
      display: flex; flex-direction: column; gap: 0;
      background: var(--surface); border: 1px solid var(--border);
      border-radius: 12px; padding: 10px; min-height: 260px;
    }
    .day-col.is-today { border-color: rgba(232,168,76,0.24); box-shadow: inset 0 0 0 1px rgba(232,168,76,0.12); }
    .day-col-header {
      font-family: var(--mono); font-size: 10px; letter-spacing: 0.07em;
      color: var(--text3); text-transform: uppercase; padding: 0 0 8px 2px;
    }
    .day-col.is-today .day-col-header { color: var(--amber); }

    .day-items { display: flex; flex-direction: column; gap: 6px; flex: 1; }
    .day-item {
      display: flex; flex-direction: column; align-items: flex-start;
      padding: 8px 9px; border-radius: 8px;
      background: var(--surface2); border: 1px solid var(--border2);
      font-size: 13px; color: var(--text); line-height: 1.4; gap: 4px;
      cursor: pointer;
      transition: border-color 0.12s, background 0.12s;
    }
    .day-item:hover { border-color: var(--amber-dim); background: rgba(232,168,76,0.06); }
    .day-col.is-today .day-item { border-color: rgba(232,168,76,0.2); }
    .day-item-row { display: flex; align-items: flex-start; justify-content: space-between; width: 100%; gap: 6px; }
    .day-item-time {
      font-family: var(--mono); font-size: 9px; color: var(--amber);
      background: var(--amber-faint); border: 1px solid var(--amber-dim);
      border-radius: 3px; padding: 1px 5px; letter-spacing: 0.02em;
      white-space: nowrap; flex-shrink: 0;
    }
    .day-item-text { flex: 1; min-width: 0; overflow: hidden; text-overflow: ellipsis; display: -webkit-box; -webkit-line-clamp: 2; -webkit-box-orient: vertical; font-size: 12px; line-height: 1.45; }
    .day-item-del {
      background: none; border: none; color: transparent; cursor: pointer;
      font-size: 14px; padding: 1px 3px; border-radius: 3px; flex-shrink: 0; transition: color 0.1s;
    }
    .day-item:hover .day-item-del { color: var(--text3); }
    .day-item-del:hover { color: #c87070 !important; }

    .day-empty {
      padding: 10px 9px; border-radius: 7px; border: 1px dashed var(--border);
      color: var(--text3); font-family: var(--mono); font-size: 10px;
      text-align: center; letter-spacing: 0.04em;
    }
    .hint {
      margin-top: 24px; font-family: var(--mono); font-size: 11px;
      color: var(--text3); letter-spacing: 0.04em; line-height: 1.9;
    }

    /* time slot row */
    .time-slot-row {
      display: flex; gap: 6px; align-items: center; margin: -20px 0 24px;
    }
    .time-slot-label {
      font-family: var(--mono); font-size: 10px; color: var(--text3);
      letter-spacing: 0.05em; flex-shrink: 0;
    }
    .time-inp {
      background: var(--surface); border: 1px solid var(--border2);
      border-radius: 6px; padding: 5px 8px;
      font-family: var(--mono); font-size: 12px; color: var(--text);
      outline: none; width: 96px;
    }
    .time-inp:focus { border-color: var(--amber-dim); }
    .time-inp::-webkit-calendar-picker-indicator { filter: invert(0.5); cursor: pointer; }
    .time-sep { font-family: var(--mono); font-size: 11px; color: var(--text3); }
    .time-clear { background: none; border: none; color: var(--text3); font-size: 13px; cursor: pointer; }
    .time-clear:hover { color: #c87070; }

    /* edit modal overlay */
    .w-modal-overlay {
      position: fixed; inset: 0; background: rgba(0,0,0,0.65);
      display: flex; align-items: center; justify-content: center;
      z-index: 999; opacity: 0; pointer-events: none; transition: opacity 0.15s;
    }
    .w-modal-overlay.open { opacity: 1; pointer-events: all; }
    .w-modal {
      background: var(--surface); border: 1px solid var(--border2); border-radius: 12px;
      padding: 22px 24px; width: 340px; max-width: calc(100vw - 48px);
    }
    .w-modal-title { font-size: 14px; color: var(--text); margin-bottom: 16px; }
    .w-modal-actions { display: flex; gap: 8px; margin-top: 16px; }
    .w-modal-cancel {
      flex: 1; padding: 8px; border-radius: 7px; border: 1px solid var(--border2);
      background: transparent; color: var(--text2); font-family: var(--sans); font-size: 13px; cursor: pointer;
    }
    .w-modal-save {
      flex: 2; padding: 8px; border-radius: 7px; border: none;
      background: var(--amber); color: #111110; font-family: var(--sans); font-size: 13px; font-weight: 500; cursor: pointer;
    }
  </style>
</head>
<body>
<div class="page">
  <div class="page-head">
    <div>
      <div class="page-title">Routine</div>
      <div class="page-sub">Weekly Planner</div>
    </div>
    <div class="page-pill">Weekly planner</div>
  </div>
  <div class="add-bar">
    <input class="add-inp" id="w-inp" placeholder="Pick a day & Add a task,👉"/>
    <div class="day-tabs" id="day-tabs"></div>
    <button class="add-go" onclick="addItem()">+ ADD</button>
  </div>

  <div class="time-slot-row">
    <span class="time-slot-label">TIME SLOT</span>
    <input class="time-inp" type="time" id="w-ts-start" title="Start time">
    <span class="time-sep">–</span>
    <input class="time-inp" type="time" id="w-ts-end" title="End time">
    <button class="time-clear" title="Clear" onclick="clearWeeklyTime()">×</button>
  </div>

  <div class="planner-scroll">
    <div class="week-grid" id="week-grid"></div>
  </div>

  <div class="hint">
    Tasks added here are injected into your routine on their scheduled day.<br>
    They appear alongside your daily routine on the Home page automatically.
  </div>
</div>

<div class="w-modal-overlay" id="w-modal" onclick="wModalOverlayClick(event)">
  <div class="w-modal">
    <div class="w-modal-title" id="w-modal-title">Edit time slot</div>
    <div class="time-slot-row" style="margin:0">
      <span class="time-slot-label">TIME</span>
      <input class="time-inp" type="time" id="w-edit-start">
      <span class="time-sep">–</span>
      <input class="time-inp" type="time" id="w-edit-end">
    </div>
    <div class="w-modal-actions">
      <button class="w-modal-cancel" onclick="clearWeeklyEdit()">Clear</button>
      <button class="w-modal-cancel" onclick="closeWeeklyModal()">Cancel</button>
      <button class="w-modal-save" onclick="saveWeeklyEdit()">Save</button>
    </div>
  </div>
</div>

<script>
const DAY_NAMES  = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday'];
const DAY_SHORT  = ['MON','TUE','WED','THU','FRI','SAT','SUN'];

let _allItems = [];   // all weekly items from server
let _selDay   = -1;   // currently selected weekday tab
let _editItemId = null;

function todayDow() {
  // 0=Mon ... 6=Sun  (matches Python weekday())
  return (new Date().getDay() + 6) % 7;
}

async function api(method, path, body) {
  const opts = { method, headers: {'Content-Type': 'application/json'} };
  if (body !== undefined) opts.body = JSON.stringify(body);
  const r = await fetch(path, opts);
  return r.json();
}

function fmtTime(hhmm) {
  if (!hhmm) return '';
  const [h, m] = hhmm.split(':').map(Number);
  const ap = h >= 12 ? 'PM' : 'AM';
  return `${h % 12 || 12}:${String(m).padStart(2,'0')} ${ap}`;
}

function sortByTime(arr) {
  return [...arr].sort((a, b) => {
    const at = a.time_start || null, bt = b.time_start || null;
    if (at && bt) return at < bt ? -1 : at > bt ? 1 : 0;
    if (at) return -1;
    if (bt) return 1;
    return 0;
  });
}

function readWeeklyTime() {
  const s = (document.getElementById('w-ts-start').value || '').trim();
  const e = (document.getElementById('w-ts-end').value   || '').trim();
  return { time_start: s || null, time_end: e || null };
}

function clearWeeklyTime() {
  document.getElementById('w-ts-start').value = '';
  document.getElementById('w-ts-end').value   = '';
}

function openWeeklyModal(item) {
  _editItemId = item.id;
  document.getElementById('w-modal-title').textContent = item.text;
  document.getElementById('w-edit-start').value = item.time_start || '';
  document.getElementById('w-edit-end').value   = item.time_end   || '';
  document.getElementById('w-modal').classList.add('open');
}

function closeWeeklyModal() {
  _editItemId = null;
  document.getElementById('w-modal').classList.remove('open');
}

function wModalOverlayClick(e) {
  if (e.target === document.getElementById('w-modal')) closeWeeklyModal();
}

async function saveWeeklyEdit() {
  if (!_editItemId) return;
  const s = (document.getElementById('w-edit-start').value || '').trim() || null;
  const e = (document.getElementById('w-edit-end').value   || '').trim() || null;
  try {
    const updated = await api('PATCH', '/api/weekly-routine/items/' + _editItemId, { time_start: s, time_end: e });
    const idx = _allItems.findIndex(x => x.id === _editItemId);
    if (idx !== -1) _allItems[idx] = { ..._allItems[idx], ...updated };
    closeWeeklyModal();
    renderGrid();
  } catch(e) { console.error('saveWeeklyEdit', e); }
}

async function clearWeeklyEdit() {
  if (!_editItemId) return;
  try {
    const updated = await api('PATCH', '/api/weekly-routine/items/' + _editItemId, { time_start: null, time_end: null });
    const idx = _allItems.findIndex(x => x.id === _editItemId);
    if (idx !== -1) _allItems[idx] = { ..._allItems[idx], ...updated };
    closeWeeklyModal();
    renderGrid();
  } catch(e) { console.error('clearWeeklyEdit', e); }
}

async function load() {
  try { _allItems = await api('GET', '/api/weekly-routine/all'); }
  catch(e) { _allItems = []; }
  renderGrid();
}

function renderDayTabs() {
  const tabs = document.getElementById('day-tabs');
  tabs.innerHTML = '';
  DAY_SHORT.forEach((label, i) => {
    const btn = document.createElement('button');
    btn.className = 'day-tab' + (_selDay === i ? ' sel' : '');
    btn.textContent = label;
    btn.onclick = () => { _selDay = (_selDay === i) ? -1 : i; renderDayTabs(); };
    tabs.appendChild(btn);
  });
}

function renderGrid() {
  const grid = document.getElementById('week-grid');
  const today = todayDow();
  grid.innerHTML = '';

  DAY_NAMES.forEach((name, i) => {
    const dayItems = sortByTime(_allItems.filter(x => x.weekday === i));
    const isToday  = i === today;

    const col = document.createElement('div');
    col.className = 'day-col' + (isToday ? ' is-today' : '');

    col.innerHTML = `<div class="day-col-header">${DAY_SHORT[i]}</div>`;

    const itemsDiv = document.createElement('div');
    itemsDiv.className = 'day-items';

    if (dayItems.length === 0) {
      itemsDiv.innerHTML = '<div class="day-empty">rest</div>';
    } else {
      dayItems.forEach(item => {
        const row = document.createElement('div');
        row.className = 'day-item';
        let timeHtml = '';
        if (item.time_start) {
          const end = item.time_end ? ' – ' + fmtTime(item.time_end) : '';
          timeHtml = `<span class="day-item-time">${fmtTime(item.time_start)}${end}</span>`;
        }
        row.innerHTML = `${timeHtml}<div class="day-item-row">
          <span class="day-item-text" title="${item.text}">${item.text}</span>
          <button class="day-item-del" title="Remove">\xd7</button></div>`;
        row.onclick = (ev) => {
          if (ev.target.classList.contains('day-item-del')) return;
          openWeeklyModal(item);
        };
        row.querySelector('.day-item-del').onclick = (ev) => { ev.stopPropagation(); deleteItem(item.id); };
        itemsDiv.appendChild(row);
      });
    }

    col.appendChild(itemsDiv);
    grid.appendChild(col);
  });
}

async function addItem() {
  const inp  = document.getElementById('w-inp');
  const text = inp.value.trim();
  if (!text) { inp.focus(); return; }
  if (_selDay < 0) {
    const tabs = document.getElementById('day-tabs');
    tabs.style.outline = '1px solid var(--amber-dim)';
    tabs.style.borderRadius = '6px';
    setTimeout(() => { tabs.style.outline = ''; }, 1200);
    return;
  }
  const slot = readWeeklyTime();
  try {
    const item = await api('POST', '/api/weekly-routine/items', { text, weekday: _selDay, ...slot });
    item.weekday = _selDay;
    _allItems.push(item);
    inp.value = '';
    clearWeeklyTime();
    renderGrid();
  } catch(e) { console.error('addItem', e); }
}

async function deleteItem(id) {
  _allItems = _allItems.filter(x => x.id !== id);
  renderGrid();
  try { await api('DELETE', '/api/weekly-routine/items/' + id); }
  catch(e) { console.error('deleteItem', e); }
}

window.addItem          = addItem;
window.clearWeeklyTime  = clearWeeklyTime;
window.closeWeeklyModal = closeWeeklyModal;
window.saveWeeklyEdit   = saveWeeklyEdit;
window.clearWeeklyEdit  = clearWeeklyEdit;
window.wModalOverlayClick = wModalOverlayClick;
window.__pageInit = function() {
  renderDayTabs();
  load();
};
</script>
</body>
</html>"""
